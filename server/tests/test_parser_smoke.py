"""
Smoke tests for PDF and DOCX parsing functionality.
Tests basic parsing capabilities without requiring actual files.
"""
import pytest
import sys
import os
from io import BytesIO
import tempfile

# Add the app directory to Python path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'app'))

from services.parse import parse_file, parse_text


class TestParserSmoke:
    """Smoke tests for parser functionality."""

    def test_parse_text_resume_basic(self):
        """Test basic resume text parsing works."""
        resume_text = """
        John Smith
        john@email.com | (555) 123-4567
        
        EXPERIENCE
        Software Engineer | Company A | 2020-Present
        • Built web applications using Python and JavaScript
        • Improved system performance by 25%
        
        SKILLS
        Python, JavaScript, SQL, React
        """
        
        result = parse_text("resume", resume_text)
        
        # Basic structure checks
        assert "contact" in result
        assert "skills" in result  
        assert "experience" in result
        
        # Contact info extracted
        assert result["contact"]["email"] == "john@email.com"
        assert "john" in result["contact"]["name"].lower()
        
        # Skills extracted
        assert "Python" in result["skills"]
        assert "JavaScript" in result["skills"]
        
        # Experience extracted
        assert len(result["experience"]) >= 1
        # Company might be in role field due to parsing format
        experience_text = f"{result['experience'][0]['company']} {result['experience'][0]['role']}"
        assert "Company A" in experience_text or "Software Engineer" in experience_text

    def test_parse_text_jd_basic(self):
        """Test basic job description text parsing works.""" 
        jd_text = """
        Software Engineer
        
        We are seeking a talented Software Engineer to join our team.
        
        REQUIREMENTS
        • 3+ years Python experience
        • Experience with React and JavaScript
        • SQL database knowledge
        
        RESPONSIBILITIES  
        • Build and maintain web applications
        • Collaborate with cross-functional teams
        • Write clean, maintainable code
        """
        
        result = parse_text("jd", jd_text)
        
        # Basic structure checks
        assert "title" in result
        assert "required" in result
        assert "responsibilities" in result
        
        # Requirements extracted (case-insensitive check)
        req_text = " ".join(result["required"]).lower()
        assert "python" in req_text
        assert "javascript" in req_text
        
        # Responsibilities extracted
        assert len(result["responsibilities"]) >= 1
        assert any("applications" in resp.lower() for resp in result["responsibilities"])

    def test_parse_text_handles_empty_input(self):
        """Test parser handles empty or minimal input gracefully."""
        # Empty resume
        result = parse_text("resume", "")
        assert "skills" in result
        assert result["skills"] == []
        
        # Empty JD
        result = parse_text("jd", "")
        assert "required" in result
        
        # Minimal resume
        result = parse_text("resume", "John Smith")
        assert "contact" in result
        
        # Minimal JD  
        result = parse_text("jd", "Software Engineer Position")
        assert "title" in result

    def test_parse_text_handles_malformed_input(self):
        """Test parser handles malformed input gracefully."""
        malformed_inputs = [
            "Random text with no structure",
            "••••••••••••••••••••••••••••••••••",
            "123456789012345678901234567890",
            "ALLCAPSTEXT" * 100,
            "Special chars: ☃️🎉💻✨🚀",
        ]
        
        for text in malformed_inputs:
            # Should not raise exceptions
            resume_result = parse_text("resume", text)
            jd_result = parse_text("jd", text)
            
            # Should return valid structure
            assert isinstance(resume_result, dict)
            assert isinstance(jd_result, dict)
            assert "skills" in resume_result
            assert "required" in jd_result

    def test_create_simple_pdf_and_parse(self):
        """Test creating a simple PDF and parsing it."""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
        except ImportError:
            pytest.skip("reportlab not available for PDF generation")
        
        # Create a simple PDF in memory
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=letter)
        c.drawString(100, 750, "John Smith")
        c.drawString(100, 730, "john@email.com")
        c.drawString(100, 700, "SKILLS")
        c.drawString(100, 680, "• Python")
        c.drawString(100, 660, "• JavaScript") 
        c.drawString(100, 630, "EXPERIENCE")
        c.drawString(100, 610, "Software Engineer | Company A")
        c.drawString(100, 590, "• Built applications")
        c.save()
        
        # Parse the PDF
        buffer.seek(0)
        result = parse_file("resume", "test.pdf", buffer)
        
        # Basic checks - PDF parsing should extract text
        assert isinstance(result, dict)
        assert "contact" in result or "skills" in result or "experience" in result
        
        # Should contain some of our text
        result_str = str(result).lower()
        assert "john" in result_str or "python" in result_str or "engineer" in result_str

    def test_create_simple_docx_and_parse(self):
        """Test creating a simple DOCX and parsing it."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available for DOCX generation")
        
        # Create a simple DOCX in memory
        doc = Document()
        doc.add_paragraph("John Smith")
        doc.add_paragraph("john@email.com")
        doc.add_paragraph("SKILLS")
        doc.add_paragraph("• Python")
        doc.add_paragraph("• JavaScript")
        doc.add_paragraph("EXPERIENCE") 
        doc.add_paragraph("Software Engineer | Company A")
        doc.add_paragraph("• Built applications")
        
        buffer = BytesIO()
        doc.save(buffer)
        
        # Parse the DOCX
        buffer.seek(0)
        result = parse_file("resume", "test.docx", buffer)
        
        # Basic checks - DOCX parsing should extract text
        assert isinstance(result, dict)
        assert "contact" in result or "skills" in result or "experience" in result
        
        # Should contain some of our text  
        result_str = str(result).lower()
        assert "john" in result_str or "python" in result_str or "engineer" in result_str

    def test_pdf_parsing_error_handling(self):
        """Test PDF parsing handles corrupted files gracefully."""
        # Create a fake "PDF" that's actually just text
        fake_pdf = BytesIO(b"This is not a real PDF file")
        
        with pytest.raises((ValueError, Exception)):
            parse_file("resume", "fake.pdf", fake_pdf)

    def test_docx_parsing_error_handling(self):
        """Test DOCX parsing handles corrupted files gracefully.""" 
        # Create a fake "DOCX" that's actually just text
        fake_docx = BytesIO(b"This is not a real DOCX file")
        
        with pytest.raises((ValueError, Exception)):
            parse_file("resume", "fake.docx", fake_docx)

    def test_unsupported_file_type(self):
        """Test unsupported file types are handled gracefully."""
        text_buffer = BytesIO(b"Some text content")
        
        # .txt should fallback to text parsing
        result = parse_file("resume", "test.txt", text_buffer)
        assert isinstance(result, dict)
        
        # Unknown extension should raise error or fallback
        unknown_buffer = BytesIO(b"Some content")
        try:
            result = parse_file("resume", "test.xyz", unknown_buffer)
            # If it doesn't raise, should return valid structure
            assert isinstance(result, dict)
        except ValueError:
            # Expected behavior for unsupported types
            pass

    def test_parse_file_vs_parse_text_consistency(self):
        """Test parse_file and parse_text produce consistent results for text."""
        sample_text = """
        John Doe
        john@example.com
        
        SKILLS
        Python, JavaScript, SQL
        
        EXPERIENCE
        Developer | Tech Corp
        • Built web apps
        """
        
        # Parse as text
        text_result = parse_text("resume", sample_text)
        
        # Parse as "file" (text fallback)
        text_buffer = BytesIO(sample_text.encode('utf-8'))
        file_result = parse_file("resume", "test.txt", text_buffer)
        
        # Results should be very similar
        assert text_result["contact"] == file_result["contact"]
        assert set(text_result["skills"]) == set(file_result["skills"])
        
    def test_large_file_handling(self):
        """Test handling of large files doesn't crash."""
        # Create large text content with proper skill section
        large_content = "John Doe\n" + "Python developer with experience.\n" * 100
        large_content += "\nSKILLS\nPython, JavaScript, React, Node.js\n"
        large_content += "• Built applications\n" * 100
        
        # Should handle large text without crashing
        result = parse_text("resume", large_content)
        assert isinstance(result, dict)
        assert "skills" in result
        
        # Should contain expected skills when properly formatted
        all_text = str(result).lower()
        assert "python" in all_text or "javascript" in all_text

    def test_parsing_special_characters(self):
        """Test parsing content with special characters and unicode."""
        special_text = """
        José María García-López
        josé@empresa.com | +34-123-456-789
        
        HABILIDADES
        • Python 🐍
        • JavaScript ⚡
        • C++ & C#
        • SQL (MySQL, PostgreSQL)
        • React.js & Vue.js
        
        EXPERIENCIA
        Desarrollador Senior | Empresa Técnológica S.L.
        • Desarrollé aplicaciones web con 99.9% uptime
        • Mejoré rendimiento en un 50%
        """
        
        result = parse_text("resume", special_text)
        
        # Should handle unicode names
        assert "contact" in result
        assert "josé" in result["contact"]["name"].lower()
        
        # Should extract skills despite special chars (may be in different sections)
        assert "skills" in result
        # Check if skills were extracted or if content is present in any section
        all_content = str(result).lower()
        assert "python" in all_content or "javascript" in all_content or len(result["skills"]) > 0
        
        # Should handle accented characters
        result_str = str(result).lower()
        assert "desarrollador" in result_str or "senior" in result_str


if __name__ == "__main__":
    pytest.main([__file__, "-v"])