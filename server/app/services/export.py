from io import BytesIO
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def _set_margins(doc: Document):
    """Set professional margins (1 inch all around)"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def _add_heading(doc: Document, text: str, size: int = 12, space_before: bool = True):
    """Add a professional section heading"""
    if space_before:
        # Add space before heading (except first one)
        spacer = doc.add_paragraph()
        spacer.space_after = Pt(6)
    
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = "Calibri"

def _add_role_header(doc: Document, role_data: Dict[str, Any]):
    """Add role/company header with professional formatting"""
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    
    # Role title (bold)
    role_run = p.add_run(role_data.get("role", ""))
    role_run.bold = True
    role_run.font.size = Pt(11)
    role_run.font.name = "Calibri"
    
    # Company (bold italic)
    company = role_data.get("company")
    if company:
        company_run = p.add_run(f" | {company}")
        company_run.bold = True
        company_run.italic = True
        company_run.font.size = Pt(11)
        company_run.font.name = "Calibri"
    
    # Dates and location on same line (right-aligned if possible, or new line)
    date_location_parts = []
    start = role_data.get("start", "")
    end = role_data.get("end", "Present")
    if start:
        date_location_parts.append(f"{start} – {end}")
    
    location = role_data.get("location")
    if location:
        date_location_parts.append(location)
    
    if date_location_parts:
        dates_run = p.add_run(f" | {' | '.join(date_location_parts)}")
        dates_run.font.size = Pt(11)
        dates_run.font.name = "Calibri"

def _add_bullets(doc: Document, bullets: List[str]):
    """Add bullet points with proper spacing and formatting"""
    for bullet in bullets:
        p = doc.add_paragraph()
        p.style = "List Bullet"
        p.space_after = Pt(2)
        
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        
        # Clear existing text and add our bullet
        p.clear()
        p.add_run(bullet)
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.name = "Calibri"

def resume_to_docx(resume: Dict[str, Any]) -> bytes:
    doc = Document()
    _set_margins(doc)
    
    # Name / Contact Header (professional formatting)
    contact = (resume or {}).get("contact") or {}
    name = contact.get("name")
    
    if name:
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = header.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.name = "Calibri"
    
    # Contact info (centered, smaller)
    contact_parts = []
    if contact.get("email"):
        contact_parts.append(contact["email"])
    if contact.get("phone"):
        contact_parts.append(contact["phone"])
    if contact.get("links"):
        contact_parts.extend(contact["links"])
    
    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.space_after = Pt(12)
        contact_run = contact_p.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(11)
        contact_run.font.name = "Calibri"

    # Summary
    if resume.get("summary"):
        _add_heading(doc, "Summary", space_before=False)
        summary_p = doc.add_paragraph(resume["summary"])
        summary_p.space_after = Pt(8)
        for run in summary_p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    # Skills
    skills = resume.get("skills") or []
    if skills:
        _add_heading(doc, "Technical Skills")
        skills_p = doc.add_paragraph(", ".join(skills))
        skills_p.space_after = Pt(8)
        for run in skills_p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    # Experience
    exp = resume.get("experience") or []
    if exp:
        _add_heading(doc, "Professional Experience")
        for i, role in enumerate(exp):
            _add_role_header(doc, role)
            _add_bullets(doc, role.get("bullets") or [])
            
            # Add small space between roles (except last)
            if i < len(exp) - 1:
                spacer = doc.add_paragraph()
                spacer.space_after = Pt(4)

    # Projects
    projects = resume.get("projects") or []
    if projects:
        _add_heading(doc, "Projects")
        for i, project in enumerate(projects):
            proj_p = doc.add_paragraph()
            proj_p.space_after = Pt(2)
            proj_run = proj_p.add_run(project.get("name") or "Project")
            proj_run.bold = True
            proj_run.font.size = Pt(11)
            proj_run.font.name = "Calibri"
            
            _add_bullets(doc, project.get("bullets") or [])
            
            # Add small space between projects (except last)
            if i < len(projects) - 1:
                spacer = doc.add_paragraph()
                spacer.space_after = Pt(4)

    # Education
    edu = resume.get("education") or []
    if edu:
        _add_heading(doc, "Education")
        for education in edu:
            edu_parts = [education.get("school"), education.get("degree"), education.get("grad")]
            edu_line = " • ".join(filter(None, edu_parts))
            if edu_line:
                edu_p = doc.add_paragraph(edu_line)
                edu_p.space_after = Pt(2)
                for run in edu_p.runs:
                    run.font.size = Pt(11)
                    run.font.name = "Calibri"

    # Other sections
    other = resume.get("other_sections") or {}
    for sec_name, items in other.items():
        if not items: 
            continue
        _add_heading(doc, str(sec_name).title())
        _add_bullets(doc, items)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()